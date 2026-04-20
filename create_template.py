"""
create_template.py

Генерирует template.docx.
Правила полей подписанта:
  client.position          — должность в родительном падеже («генерального директора»)
  client.signatory         — ФИО в родительном падеже («Лаврова Антона Алексеевича»)
  client.signatory_short   — Фамилия И.О. в именительном падеже («Лавров А.А.»)
  client.authority         — основание полномочий («Устава» или «Доверенности № 6 от 04.04.2025 г.»)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CONTRACTOR = {
    "name":              'ООО «Персонально Ваш»',
    "inn":               '9710037361',
    "kpp":               '771001001',
    "address":           '123056, город Москва, ул Юлиуса Фучика, д. 6 стр. 2, помещ. 6ч',
    "signatory":         'Марушевской Виктории',
    "signatory_short":   'Марушевская В.',
    "position":          'генерального директора',
    "authority":         'Устава',
    "bank_account":      '40702810600000053430',
    "bank_name":         'АО "Райффайзенбанк", г. Москва',
    "bank_corr_account": '30101810200000000700',
    "bank_bik":          '044525700',
    "phone":             '+7 495 128 2110',
}


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(edge)
        if val:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, size=11):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def add_paragraph(doc, text='', bold=False, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align
    if text:
        add_run(p, text, bold=bold, size=size)
    return p


def make_template():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    add_paragraph(doc,
        'Акт сдачи-приёмки услуг к Счет-договору № {{invoice_number}} от {{invoice_date}}г.',
        bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    add_run(p2, 'г. Москва', size=11)
    add_run(p2, '\t', size=11)
    add_run(p2, '{{act_date}}г.', size=11)
    pPr = p2._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9072')
    tabs.append(tab_el)
    pPr.append(tabs)

    c = CONTRACTOR
    intro = (
        f'{c["name"]}, именуемое в дальнейшем «Исполнитель», '
        f'в лице {c["position"]} {c["signatory"]}, '
        f'действующей на основании {c["authority"]}, с одной стороны, '
        f'и {{{{client.name}}}}, именуемое в дальнейшем «Заказчик», '
        f'в лице {{{{client.position}}}} {{{{client.signatory}}}}, '
        f'действующего на основании {{{{client.authority}}}}, с другой стороны, '
        f'вместе и по отдельности именуемые «Стороны», составили настоящий акт '
        f'сдачи-приёмки услуг к Счет-договору №\u00a0{{{{invoice_number}}}} '
        f'от\u00a0{{{{invoice_date}}}}г. (далее\u00a0–\u00a0Акт):'
    )
    add_paragraph(doc, intro, size=11, space_after=8)

    add_paragraph(doc,
        '1. Исполнитель оказал {{service_name}} – {{hours}} ({{hours_text}}) часов.',
        size=11, space_after=6)
    add_paragraph(doc,
        '2. Сумма оказанных услуг составила {{total_amount}} ({{total_amount_text}} рублей), 00\u00a0коп. '
        'НДС не облагается на основании применения Исполнителем упрощенной системы '
        'налогообложения в соответствии со ст.\u00a0346.12, 346.13 Главы\u00a026.2 '
        'Налогового кодекса Российской Федерации.',
        size=11, space_after=6)
    add_paragraph(doc, '3. Материалы сессий доступны по ссылке {{payment_link}}', size=11, space_after=6)
    add_paragraph(doc,
        '4. Консалтинговые услуги по Счет-договору №\u00a0{{invoice_number}} '
        'от\u00a0{{invoice_date}}г. выполнены полностью и в срок. '
        'Претензий по объему, качеству результата работ и срокам их выполнения Заказчик не имеет.',
        size=11, space_after=6)
    add_paragraph(doc, '5. Акт составлен в двух экземплярах, по одному для каждой из сторон.', size=11, space_after=14)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Cm(8.5)
    tbl.columns[1].width = Cm(8.5)

    cell_c = tbl.rows[0].cells[0]
    cell_c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for i, (text, bold) in enumerate([
        ('Исполнитель:', True),
        (c['name'], False), (f'ИНН {c["inn"]}', False), (f'КПП {c["kpp"]}', False),
        (c['address'], False), ('', False),
        ('Банковские реквизиты:', True),
        (f'р/с {c["bank_account"]}', False), (c['bank_name'], False),
        (f'к/с {c["bank_corr_account"]}', False), (f'БИК {c["bank_bik"]}', False),
        (f'Тел. {c["phone"]}', False),
    ]):
        p = cell_c.paragraphs[0] if i == 0 else cell_c.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, text, bold=bold, size=10)

    cell_cl = tbl.rows[0].cells[1]
    cell_cl.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for i, (text, bold) in enumerate([
        ('Заказчик:', True),
        ('{{client.name}}', False), ('ИНН {{client.inn}}', False), ('КПП {{client.kpp}}', False),
        ('{{client.address}}', False), ('', False),
        ('Банковские реквизиты:', True),
        ('р/с {{client.bank_account}}', False), ('{{client.bank_name}}', False),
        ('к/с {{client.bank_corr_account}}', False), ('БИК {{client.bank_bik}}', False),
        ('Тел. {{client.phone}}', False),
    ]):
        p = cell_cl.paragraphs[0] if i == 0 else cell_cl.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, text, bold=bold, size=10)

    doc.add_paragraph()

    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.autofit = False
    sig_tbl.columns[0].width = Cm(8.5)
    sig_tbl.columns[1].width = Cm(8.5)

    def sig_cell(cell, text, bold=False):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, text, bold=bold, size=11)

    sig_cell(sig_tbl.rows[0].cells[0], 'Исполнитель', bold=True)
    sig_cell(sig_tbl.rows[0].cells[1], 'Заказчик', bold=True)
    sig_cell(sig_tbl.rows[1].cells[0], c['position'].capitalize())
    sig_cell(sig_tbl.rows[1].cells[1], '{{client.position}}')
    sig_cell(sig_tbl.rows[2].cells[0], f'Подпись _________________ / {c["signatory_short"]} /  М.П.')
    sig_cell(sig_tbl.rows[2].cells[1], 'Подпись _________________ / {{client.signatory_short}} /  М.П.')

    for row in sig_tbl.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'none'}, bottom={'val': 'none'},
                            left={'val': 'none'}, right={'val': 'none'})

    doc.save('template.docx')
    print('✅ template.docx создан')


if __name__ == '__main__':
    make_template()
