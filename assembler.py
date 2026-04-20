"""
assembler.py

Сборка итогового docx-акта из шаблона.
QR-код не используется — ссылка передаётся текстом в пункте 3.

Правила полей подписанта:
  signatory        — родительный падеж, полное ФИО (для вводного абзаца)
                     пример: «Мартынова Дмитрия Сергеевича»
  signatory_short  — именительный падеж, Фамилия И.О. (для строки подписи)
                     пример: «Мартынов Д.С.»
  Оба поля передаются явно из data — автогенерация не используется.

Правило имени файла:
  «Акт к Счет-договору № {invoice_number} от {invoice_date}г.docx»
"""

from __future__ import annotations

import copy
from pathlib import Path
from typing import Any

from docx import Document


# ---------------------------------------------------------------------------
# Вспомогательные функции
# ---------------------------------------------------------------------------

def _flatten(data: dict[str, Any], prefix: str = "") -> dict[str, str]:
    """Разворачивает вложенный dict в плоский: {"contractor.name": "..."}."""
    result: dict[str, str] = {}
    for k, v in data.items():
        full_key = f"{prefix}{k}" if prefix else k
        if isinstance(v, dict):
            result.update(_flatten(v, prefix=full_key + "."))
        elif v is None:
            result[full_key] = ""
        else:
            result[full_key] = str(v)
    return result


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Заменяет {{KEY}} в параграфе, сохраняя форматирование первого рана."""
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", value)
    if new_text != full_text:
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""


def _make_filename(data: dict[str, Any]) -> str:
    """
    Генерирует имя файла по правилу:
    «Акт к Счет-договору № {invoice_number} от {invoice_date}г.docx»
    """
    num = data.get("invoice_number", "")
    date = data.get("invoice_date", "")
    return f"Акт к Счет-договору № {num} от {date}г.docx"


# ---------------------------------------------------------------------------
# Главная функция
# ---------------------------------------------------------------------------

def assemble(
    data: dict[str, Any],
    template_path: str | Path = "template.docx",
    output_path: str | Path | None = None,
) -> Path:
    """
    Собирает документ акта из шаблона и данных.

    Args:
        data:          Валидированный JSON-словарь.
                       client.signatory        — родительный падеж (шапка)
                       client.signatory_short  — именительный Фамилия И.О. (подпись)
        template_path: Путь к шаблону .docx.
        output_path:   Куда сохранить результат. Если None — имя генерируется автоматически.

    Returns:
        Path к готовому файлу.
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")

    if output_path is None:
        output_path = Path(_make_filename(data))
    output_path = Path(output_path)

    data = copy.deepcopy(data)
    flat = _flatten(data)

    doc = Document(str(template_path))

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        _replace_in_paragraph(para, flat)

    doc.save(str(output_path))
    return output_path
